#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
УНИВЕРСАЛЬНЫЙ ГЕНЕРАТОР ЭКЗАМЕНАЦИОННЫХ БИЛЕТОВ
Автор: Темур Тураев (temurturayev7822@gmail.com)
Версия: 1.0
Дата: 2024

Функции:
- Извлечение вопросов из Word документов (.docx)
- Создание экзаменационных билетов с заданным количеством
- Поддержка разных языков (узбекский, русский)
- Автоматическое повторение вопросов при их нехватке
- Красивое форматирование выходных документов
"""

import random
import re
from docx import Document
from typing import List, Dict, Tuple
import os
import argparse
from datetime import datetime

class TicketGenerator:
    """Класс для генерации экзаменационных билетов"""
    
    def __init__(self, base_path: str):
        """
        Инициализация генератора
        
        Args:
            base_path: Путь к папке с файлами вопросов
        """
        self.base_path = base_path
        self.questions = {}
        
    def extract_questions_from_file(self, file_path: str) -> List[Tuple[int, str]]:
        """
        Извлекает вопросы из Word файла
        
        Args:
            file_path: Путь к файлу с вопросами
            
        Returns:
            Список кортежей (номер_вопроса, текст_вопроса)
        """
        try:
            doc = Document(file_path)
            questions = []
            
            if not doc.tables:
                print(f"⚠️  В файле {os.path.basename(file_path)} нет таблиц")
                return questions
            
            # Анализируем первую таблицу
            table = doc.tables[0]
            
            print(f"🔍 Анализ файла {os.path.basename(file_path)}: {len(table.rows)} строк, {len(table.columns)} столбцов")
            
            question_number = 1
            
            for i, row in enumerate(table.rows):
                if i == 0:  # Пропускаем заголовок
                    continue
                
                cells = row.cells
                
                if len(cells) >= 2:
                    # Стандартная структура: номер | вопрос
                    number_cell = cells[0].text.strip()
                    question_cell = cells[1].text.strip()
                    
                    if question_cell and len(question_cell) > 10:
                        # Пытаемся извлечь номер
                        number_match = re.search(r'\d+', number_cell)
                        if number_match:
                            num = int(number_match.group())
                        else:
                            num = question_number
                        
                        questions.append((num, question_cell))
                        question_number += 1
                else:
                    # Если только одна колонка, берем весь текст
                    full_text = ' '.join([cell.text.strip() for cell in cells if cell.text.strip()])
                    if len(full_text) > 15:
                        # Убираем начальный номер если есть
                        clean_text = re.sub(r'^\d+[\.\)]\s*', '', full_text)
                        if len(clean_text) > 10:
                            questions.append((question_number, clean_text))
                            question_number += 1
            
            print(f"📊 Извлечено {len(questions)} вопросов из {os.path.basename(file_path)}")
            
            return questions
            
        except Exception as e:
            print(f"❌ Ошибка при чтении файла {file_path}: {e}")
            return []
    
    def prepare_questions(self, questions: List[Tuple[int, str]], needed: int, subject_name: str) -> List[Tuple[int, str]]:
        """
        Подготавливает список вопросов нужного размера
        
        Args:
            questions: Список исходных вопросов
            needed: Требуемое количество вопросов
            subject_name: Название предмета для логирования
            
        Returns:
            Подготовленный список вопросов
        """
        if len(questions) >= needed:
            shuffled = questions.copy()
            random.shuffle(shuffled)
            result = shuffled[:needed]
            print(f"   ✅ {subject_name}: используем {needed} из {len(questions)} вопросов (без повторов)")
            return result
        else:
            result = []
            full_cycles = needed // len(questions)
            remainder = needed % len(questions)
            
            print(f"   ⚠️  {subject_name}: {len(questions)} вопросов × {full_cycles} циклов + {remainder} дополнительных")
            
            # Добавляем полные циклы
            for cycle in range(full_cycles):
                shuffled_cycle = questions.copy()
                random.shuffle(shuffled_cycle)
                result.extend(shuffled_cycle)
            
            # Добавляем остаток
            if remainder > 0:
                shuffled_remainder = questions.copy()
                random.shuffle(shuffled_remainder)
                result.extend(shuffled_remainder[:remainder])
            
            return result
    
    def create_tickets(self, subject_files: Dict[str, str], subject_names: Dict[str, str], 
                      num_tickets: int) -> List[Dict]:
        """
        Создает экзаменационные билеты
        
        Args:
            subject_files: Словарь {ключ_предмета: путь_к_файлу}
            subject_names: Словарь {ключ_предмета: название_для_отображения}
            num_tickets: Количество билетов
            
        Returns:
            Список билетов
        """
        # Извлекаем вопросы из всех файлов
        all_questions = {}
        for key, file_path in subject_files.items():
            questions = self.extract_questions_from_file(file_path)
            if not questions:
                raise ValueError(f"Не удалось извлечь вопросы из файла: {file_path}")
            all_questions[key] = questions
        
        print(f"\n🎯 Создаем {num_tickets} билетов из:")
        for key, questions in all_questions.items():
            print(f"   {subject_names[key]}: {len(questions)} вопросов")
        
        # Подготавливаем вопросы для каждого предмета
        prepared_questions = {}
        for key, questions in all_questions.items():
            prepared_questions[key] = self.prepare_questions(
                questions, num_tickets, subject_names[key]
            )
        
        # Создаем билеты
        tickets = []
        for i in range(num_tickets):
            ticket = {'number': i + 1}
            
            for key in subject_files.keys():
                ticket[key] = {
                    'question': prepared_questions[key][i][1],
                    'number': prepared_questions[key][i][0]
                }
            
            tickets.append(ticket)
        
        return tickets
    
    def save_tickets(self, tickets: List[Dict], subject_names: Dict[str, str], 
                    output_path: str, title: str = "ЭКЗАМЕНАЦИОННЫЕ БИЛЕТЫ"):
        """
        Сохраняет билеты в Word документ
        
        Args:
            tickets: Список билетов
            subject_names: Названия предметов для отображения
            output_path: Путь для сохранения файла
            title: Заголовок документа
        """
        doc = Document()
        
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Настройка полей
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Заголовок
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.add_run(f'Количество билетов: {len(tickets)} • ').bold = True
        info.add_run(f'Предметы: {", ".join(subject_names.values())}').italic = True
        
        date_para = doc.add_paragraph(f'Создано: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Билеты
        subject_keys = list(subject_names.keys())
        
        for i, ticket in enumerate(tickets):
            # Заголовок билета
            ticket_header = doc.add_heading(f'БИЛЕТ № {ticket["number"]}', level=1)
            ticket_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Вопросы
            for j, key in enumerate(subject_keys):
                subj_title = f'{j+1}. {subject_names[key].upper()}'
                subj_data = ticket[key]
                
                # Заголовок предмета
                subj_para = doc.add_paragraph()
                run = subj_para.add_run(subj_title)
                run.bold = True
                run.font.size = Pt(12)
                
                # Номер вопроса
                num_para = doc.add_paragraph(f'   Вопрос №{subj_data["number"]}:')
                num_para.runs[0].italic = True
                num_para.runs[0].font.size = Pt(10)
                
                # Текст вопроса
                q_para = doc.add_paragraph(subj_data['question'])
                q_para.paragraph_format.left_indent = Inches(0.4)
                q_para.paragraph_format.right_indent = Inches(0.2)
                q_para.paragraph_format.space_after = Pt(8)
                
            # Разделитель
            if i < len(tickets) - 1:
                separator = doc.add_paragraph('─' * 70)
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Новая страница каждые 3 билета
                if (i + 1) % 3 == 0:
                    doc.add_page_break()
        
        doc.save(output_path)
        print(f"\n✅ Билеты сохранены: {output_path}")

def main():
    """Основная функция для запуска из командной строки"""
    parser = argparse.ArgumentParser(description='Генератор экзаменационных билетов')
    parser.add_argument('--path', '-p', default='.', help='Путь к папке с файлами')
    parser.add_argument('--tickets', '-t', type=int, default=360, help='Количество билетов')
    parser.add_argument('--language', '-l', choices=['uzbek', 'russian'], default='uzbek', 
                       help='Язык билетов')
    parser.add_argument('--output', '-o', help='Имя выходного файла')
    
    args = parser.parse_args()
    
    print("🎓 УНИВЕРСАЛЬНЫЙ ГЕНЕРАТОР ЭКЗАМЕНАЦИОННЫХ БИЛЕТОВ")
    print("=" * 60)
    
    generator = TicketGenerator(args.path)
    
    if args.language == 'uzbek':
        # Узбекские файлы
        subject_files = {
            'adabiyot': os.path.join(args.path, "Адабиёт ўзбек савол.docx"),
            'tarbiya': os.path.join(args.path, "Тарбия ўзбек савол.docx"),
            'tarikh': os.path.join(args.path, "Тарих ўзбек савол.docx")
        }
        subject_names = {
            'adabiyot': 'АДАБИЁТ',
            'tarbiya': 'ТАРБИЯ', 
            'tarikh': 'ТАРИХ'
        }
        title = "ЭКЗАМЕНАЦИОННЫЕ БИЛЕТЫ"
        default_output = f"УЗБЕКСКИЕ_БИЛЕТЫ_{args.tickets}.docx"
    else:
        # Русские файлы
        subject_files = {
            'literatura': os.path.join(args.path, "Адабиёт рус савол.docx"),
            'vospitanie': os.path.join(args.path, "Тарбия рус савол.docx"),
            'istoriya': os.path.join(args.path, "Тарих рус савол.docx")
        }
        subject_names = {
            'literatura': 'ЛИТЕРАТУРА',
            'vospitanie': 'ВОСПИТАНИЕ',
            'istoriya': 'ИСТОРИЯ'
        }
        title = "ЭКЗАМЕНАЦИОННЫЕ БИЛЕТЫ\n(для русскоговорящих студентов)"
        default_output = f"РУССКИЕ_БИЛЕТЫ_{args.tickets}.docx"
    
    output_path = args.output or os.path.join(args.path, default_output)
    
    try:
        # Проверяем наличие файлов
        missing_files = []
        for key, file_path in subject_files.items():
            if not os.path.exists(file_path):
                missing_files.append(os.path.basename(file_path))
        
        if missing_files:
            print(f"❌ Не найдены файлы: {', '.join(missing_files)}")
            return
        
        # Создаем билеты
        print(f"\n🔍 Извлекаем вопросы...")
        tickets = generator.create_tickets(subject_files, subject_names, args.tickets)
        
        # Сохраняем
        generator.save_tickets(tickets, subject_names, output_path, title)
        
        print(f"\n🎉 УСПЕШНО ЗАВЕРШЕНО!")
        print(f"📁 Файл: {os.path.basename(output_path)}")
        print(f"📊 Билетов: {len(tickets)}")
        
    except Exception as e:
        print(f"\n❌ ОШИБКА: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    # Если запускается без аргументов, используем интерактивный режим
    import sys
    if len(sys.argv) == 1:
        print("🎓 УНИВЕРСАЛЬНЫЙ ГЕНЕРАТОР ЭКЗАМЕНАЦИОННЫХ БИЛЕТОВ")
        print("=" * 60)
        print("\nИнтерактивный режим")
        print("Для помощи запустите: python генератор.py --help")
        
        # Простой интерактивный режим
        path = input("\nПуть к папке с файлами (Enter для текущей): ").strip() or "."
        
        print("\nВыберите язык:")
        print("1. Узбекский (360 билетов)")
        print("2. Русский (160 билетов)")
        choice = input("Ваш выбор (1/2): ").strip()
        
        if choice == "2":
            sys.argv = [sys.argv[0], "--path", path, "--language", "russian", "--tickets", "160"]
        else:
            sys.argv = [sys.argv[0], "--path", path, "--language", "uzbek", "--tickets", "360"]
    
    main()